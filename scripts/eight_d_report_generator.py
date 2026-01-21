#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
8D报告自动生成器
根据收集的信息自动生成专业的8D分析报告
"""

import json
import os
from datetime import datetime, timedelta
from typing import Dict, List, Optional
from dataclasses import dataclass, asdict

@dataclass
class D0Data:
    """D0阶段数据：问题发现和初步响应"""
    problem_description: str
    discovery_date: str
    discovery_person: str
    affected_products: List[str]
    initial_severity: str
    initial_response: str

@dataclass
class D1Data:
    """D1阶段数据：组建跨功能团队"""
    team_leader: str
    team_members: List[str]
    team_roles: Dict[str, str]
    communication_plan: str

@dataclass
class D2Data:
    """D2阶段数据：问题定义和描述"""
    problem_statement: str
    problem_scope: str
    affected_customers: int
    customer_impact: str
    safety_impact: str
    legal_impact: str
    financial_impact: str

@dataclass
class D3Data:
    """D3阶段数据：临时遏制措施"""
    containment_actions: List[str]
    implementation_date: str
    responsible_person: str
    effectiveness_verification: str
    customer_notification: bool

@dataclass
class D4Data:
    """D4阶段数据：根因分析"""
    root_cause_analysis: str
    fishbone_diagram: Dict
    five_whys: List[str]
    data_analysis: Dict
    potential_causes: List[str]
    verified_root_cause: str

@dataclass
class D5Data:
    """D5阶段数据：永久纠正措施"""
    corrective_actions: List[Dict]
    implementation_plan: str
    responsible_person: str
    target_date: str
    resource_requirements: str

@dataclass
class D6Data:
    """D6阶段数据：实施和验证纠正措施"""
    implementation_status: str
    verification_results: Dict
    effectiveness_assessment: str
    side_effects: str

@dataclass
class D7Data:
    """D7阶段数据：预防再发生"""
    prevention_measures: List[str]
    process_improvements: List[str]
    training_requirements: List[str]
    system_updates: List[str]
    documentation_changes: List[str]

@dataclass
class D8Data:
    """D8阶段数据：团队总结和认可"""
    lessons_learned: List[str]
    team_recognition: str
    process_improvements: str
    knowledge_sharing: str

class EightDReportGenerator:
    """8D报告生成器"""

    def __init__(self):
        self.report_template = self._load_template()
        self.current_data = {}

    def _load_template(self) -> Dict:
        """加载8D报告模板"""
        return {
            "D0": {
                "title": "D0：问题发现和初步响应",
                "description": "识别和定义问题，建立初步响应",
                "required_fields": [
                    "problem_description",
                    "discovery_date",
                    "discovery_person",
                    "affected_products",
                    "initial_severity",
                    "initial_response"
                ]
            },
            "D1": {
                "title": "D1：组建跨功能团队",
                "description": "组建具备解决问题技能的跨功能团队",
                "required_fields": [
                    "team_leader",
                    "team_members",
                    "team_roles",
                    "communication_plan"
                ]
            },
            "D2": {
                "title": "D2：问题定义和描述",
                "description": "明确问题的详细定义和范围",
                "required_fields": [
                    "problem_statement",
                    "problem_scope",
                    "affected_customers",
                    "customer_impact",
                    "safety_impact",
                    "legal_impact",
                    "financial_impact"
                ]
            },
            "D3": {
                "title": "D3：临时遏制措施",
                "description": "实施临时措施以遏制问题影响",
                "required_fields": [
                    "containment_actions",
                    "implementation_date",
                    "responsible_person",
                    "effectiveness_verification",
                    "customer_notification"
                ]
            },
            "D4": {
                "title": "D4：根因分析",
                "description": "使用分析工具识别问题的根本原因",
                "required_fields": [
                    "root_cause_analysis",
                    "fishbone_diagram",
                    "five_whys",
                    "data_analysis",
                    "potential_causes",
                    "verified_root_cause"
                ]
            },
            "D5": {
                "title": "D5：永久纠正措施",
                "description": "制定和选择永久性纠正措施",
                "required_fields": [
                    "corrective_actions",
                    "implementation_plan",
                    "responsible_person",
                    "target_date",
                    "resource_requirements"
                ]
            },
            "D6": {
                "title": "D6：实施和验证纠正措施",
                "description": "实施纠正措施并验证其有效性",
                "required_fields": [
                    "implementation_status",
                    "verification_results",
                    "effectiveness_assessment",
                    "side_effects"
                ]
            },
            "D7": {
                "title": "D7：预防再发生",
                "description": "修改系统、程序、流程以预防问题再发生",
                "required_fields": [
                    "prevention_measures",
                    "process_improvements",
                    "training_requirements",
                    "system_updates",
                    "documentation_changes"
                ]
            },
            "D8": {
                "title": "D8：团队总结和认可",
                "description": "总结经验教训，认可团队贡献",
                "required_fields": [
                    "lessons_learned",
                    "team_recognition",
                    "process_improvements",
                    "knowledge_sharing"
                ]
            }
        }

    def collect_information(self, phase: str, user_input: Dict) -> bool:
        """
        收集特定阶段的信息

        Args:
            phase: 8D阶段 (D0-D8)
            user_input: 用户输入的信息

        Returns:
            bool: 是否成功收集信息
        """
        try:
            if phase not in self.report_template:
                raise ValueError(f"无效的阶段: {phase}")

            required_fields = self.report_template[phase]["required_fields"]

            # 检查必要字段
            missing_fields = [field for field in required_fields if field not in user_input]

            if missing_fields:
                return False

            # 根据阶段存储数据
            if phase == "D0":
                self.current_data["D0"] = D0Data(**user_input)
            elif phase == "D1":
                self.current_data["D1"] = D1Data(**user_input)
            elif phase == "D2":
                self.current_data["D2"] = D2Data(**user_input)
            elif phase == "D3":
                self.current_data["D3"] = D3Data(**user_input)
            elif phase == "D4":
                self.current_data["D4"] = D4Data(**user_input)
            elif phase == "D5":
                self.current_data["D5"] = D5Data(**user_input)
            elif phase == "D6":
                self.current_data["D6"] = D6Data(**user_input)
            elif phase == "D7":
                self.current_data["D7"] = D7Data(**user_input)
            elif phase == "D8":
                self.current_data["D8"] = D8Data(**user_input)

            return True

        except Exception as e:
            print(f"信息收集错误: {e}")
            return False

    def generate_five_whys(self, problem_statement: str) -> List[str]:
        """
        生成5Why分析

        Args:
            problem_statement: 问题描述

        Returns:
            List[str]: 5Why分析结果
        """
        # 这里可以实现更智能的5Why分析逻辑
        # 现在提供模板
        whys = [
            f"为什么 {problem_statement}？",
            "为什么会出现这种情况？",
            "为什么之前的控制措施没有生效？",
            "为什么系统没有检测到这个问题？",
            "为什么预防措施不够完善？"
        ]
        return whys

    def generate_fishbone_diagram(self, problem: str) -> Dict:
        """
        生成鱼骨图分析

        Args:
            problem: 问题描述

        Returns:
            Dict: 鱼骨图数据
        """
        return {
            "问题": problem,
            "人员": {
                "原因1": "培训不足",
                "原因2": "技能水平不够",
                "原因3": "沟通不良"
            },
            "机器": {
                "原因1": "设备老化",
                "原因2": "维护不当",
                "原因3": "校准偏差"
            },
            "材料": {
                "原因1": "供应商质量问题",
                "原因2": "存储条件不当",
                "原因3": "规格不符"
            },
            "方法": {
                "原因1": "作业标准不清",
                "原因2": "流程不合理",
                "原因3": "检查点缺失"
            },
            "环境": {
                "原因1": "温湿度异常",
                "原因2": "清洁度不够",
                "原因3": "空间限制"
            }
        }

    def check_completion_status(self) -> Dict:
        """检查8D报告完成状态"""
        status = {}
        total_phases = 8

        for i in range(total_phases):
            phase_key = f"D{i}"
            if phase_key in self.current_data:
                status[phase_key] = "已完成"
            else:
                status[phase_key] = "待完成"

        status["完成度"] = f"{len(self.current_data)}/{total_phases}"
        return status

    def generate_report(self, output_path: str = "8D_report.json") -> str:
        """
        生成完整的8D报告

        Args:
            output_path: 输出文件路径

        Returns:
            str: 报告文件路径
        """
        report_data = {
            "报告信息": {
                "生成时间": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "版本": "1.0",
                "状态": "进行中"
            },
            "8D分析": {}
        }

        # 转换数据并添加到报告中
        for phase, data in self.current_data.items():
            report_data["8D分析"][phase] = asdict(data)

        # 添加完成状态
        completion_status = self.check_completion_status()
        report_data["完成状态"] = completion_status

        # 保存报告
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(report_data, f, ensure_ascii=False, indent=2)

        return output_path

    def export_to_word(self, output_path: str = "8D_report.docx") -> str:
        """
        导出为Word文档格式（需要python-docx库）

        Args:
            output_path: 输出文件路径

        Returns:
            str: 输出文件路径
        """
        try:
            from docx import Document
            from docx.shared import Inches

            doc = Document()

            # 添加标题
            title = doc.add_heading('8D问题解决报告', 0)

            # 添加报告信息
            doc.add_heading('报告信息', level=1)
            info_table = doc.add_table(rows=3, cols=2)
            info_table.style = 'Table Grid'
            info_table.cell(0, 0).text = '生成时间'
            info_table.cell(0, 1).text = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            info_table.cell(1, 0).text = '版本'
            info_table.cell(1, 1).text = '1.0'
            info_table.cell(2, 0).text = '状态'
            info_table.cell(2, 1).text = '进行中'

            # 添加8D内容
            for phase, data in self.current_data.items():
                doc.add_heading(f'{phase}阶段', level=1)
                if hasattr(data, '__dict__'):
                    for field, value in asdict(data).items():
                        if isinstance(value, (list, dict)):
                            p = doc.add_paragraph()
                            p.add_run(f'{field}: ').bold = True
                            p.add_run(str(value))
                        else:
                            p = doc.add_paragraph()
                            p.add_run(f'{field}: ').bold = True
                            p.add_run(str(value))

            doc.save(output_path)
            return output_path

        except ImportError:
            print("需要安装python-docx库来生成Word文档")
            return ""

def main():
    """主函数，用于测试"""
    generator = EightDReportGenerator()

    # 测试D0阶段
    d0_data = {
        "problem_description": "空调制冷效果不佳",
        "discovery_date": "2024-01-20",
        "discovery_person": "张三",
        "affected_products": ["AC-2024-001", "AC-2024-002"],
        "initial_severity": "中等",
        "initial_response": "暂停相关产品出货"
    }

    if generator.collect_information("D0", d0_data):
        print("D0阶段信息收集成功")

    # 生成5Why分析
    whys = generator.generate_five_whys("空调制冷效果不佳")
    print("5Why分析:", whys)

    # 生成鱼骨图
    fishbone = generator.generate_fishbone_diagram("空调制冷效果不佳")
    print("鱼骨图:", fishbone)

    # 检查完成状态
    status = generator.check_completion_status()
    print("完成状态:", status)

    # 生成报告
    report_file = generator.generate_report("test_8D_report.json")
    print(f"报告已生成: {report_file}")

if __name__ == "__main__":
    main()
